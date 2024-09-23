# Derechos de autor 2024 Eduardo Augusto García Rodríguez
# Este script de auditoría recoge información del sistema y la guarda en un informe.

import os
import platform
import subprocess
from docx import Document

output_file = "audit_report.docx"

def get_system_info():
    """Recopila información sobre el sistema operativo y el hardware."""
    info = {}
    info['Sistema Operativo'] = platform.system()
    info['Versión del SO'] = platform.version()
    info['Plataforma'] = platform.platform()
    info['Nombre de la máquina'] = platform.node()
    info['Arquitectura'] = platform.architecture()
    return info

def get_installed_software():
    """Obtiene la lista de software instalado. Funciona en Windows."""
    if platform.system() == "Windows":
        return subprocess.getoutput("wmic product get name")
    else:
        return "Software instalado no disponible en este SO"

def generate_report():
    """Genera un informe con la información del sistema y el software instalado."""
    doc = Document()
    doc.add_heading('Informe de Auditoría del Sistema', 0)

    # Obtener información del sistema
    system_info = get_system_info()
    doc.add_heading('Información del Sistema', level=1)
    for key, value in system_info.items():
        doc.add_paragraph(f"{key}: {value}")

    # Obtener software instalado
    installed_software = get_installed_software()
    doc.add_heading('Software Instalado', level=1)
    doc.add_paragraph(installed_software)

    # Guardar el informe en el archivo
    doc.save(output_file)
    print(f"Reporte guardado en {output_file}")

if __name__ == "__main__":
    generate_report()
